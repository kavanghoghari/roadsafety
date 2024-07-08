import pytesseract
from PIL import Image
import pandas as pd
from docx import Document
import os

# Configure pytesseract to point to the tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'path_to_your_tesseract_executable'

def extract_text_from_image(image_path):
    # Open an image file
    with Image.open(image_path) as img:
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(img)
    return text

def process_images(image_folder, output_csv, output_docx):
    # Prepare a list to hold the data
    data = []

    # Create a Word document
    doc = Document()

    # Loop through all images in the folder
    for filename in os.listdir(image_folder):
        if filename.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif')):
            image_path = os.path.join(image_folder, filename)
            text = extract_text_from_image(image_path)

            # Split the text into lines
            lines = text.split('\n')

            # Extract the last line (note) and other lines
            note = lines[-1] if lines else ""
            other_info = "\n".join(lines[:-1]) if lines else ""

            # Append the extracted data to the list
            data.append([filename, other_info, note])

            # Write to Word document
            doc.add_heading(filename, level=1)
            doc.add_paragraph(other_info)
            doc.add_paragraph(note)
            doc.add_page_break()

    # Create a DataFrame from the data
    df = pd.DataFrame(data, columns=['Filename', 'Information', 'Note'])

    # Write the DataFrame to a CSV file
    df.to_csv(output_csv, index=False)

    # Save the Word document
    doc.save(output_docx)

# Folder containing the images
image_folder = '/path_tamara_image_folder_nu'

# Output CSV file
output_csv = '/tamaru_output_nu_path.csv'

# Output Word file
output_docx = '/tamaru_output_nu_path1.docx'

# Process the images and generate the CSV and Word file
process_images(image_folder, output_csv, output_docx)
